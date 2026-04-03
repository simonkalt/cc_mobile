"""
Generate a .docx (Word) document from cover letter content (markdown or HTML).

Preserves bold, italic, and lists in the output. Used so the server can return
a fully adorned .docx to the client for editing; PDF export is handled by the
Syncfusion (.NET) service (not POST /api/files/docx-to-pdf on this API).

When USE_DOCX_COMPONENTS is True, the LLM may return three XML components
(document_xml, numbering_xml, styles_xml) which we assemble into a .docx via
build_docx_from_components().
"""

import io
import os
import zipfile
import html as htmllib
import logging
import re
from html.parser import HTMLParser
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    RGBColor = None
    Inches = None
    WD_BREAK = None
    OxmlElement = None
    qn = None

_VISUAL_BULLET_RE = re.compile(r"^\s*[•◦▪▸\-\*\+]\s+")
_VISUAL_NUMBER_RE = re.compile(r"^\s*(?:\(?[1-9]\d?\)?[.)])\s+")


def _apply_visual_hanging_indent_fallback(doc) -> None:
    """
    Enforce hanging indent on paragraphs that look like list items by text prefix.
    This is a safety net when upstream parsing or model XML leaves visible bullets/numbers.
    """
    if not DOCX_AVAILABLE or Inches is None:
        return
    list_left = Inches(0.25)
    list_hang = Inches(-0.25)
    try:
        for p in doc.paragraphs:
            txt = (p.text or "").strip()
            if not txt:
                continue
            if _VISUAL_BULLET_RE.match(txt) or _VISUAL_NUMBER_RE.match(txt):
                p.paragraph_format.left_indent = list_left
                p.paragraph_format.first_line_indent = list_hang
    except Exception as e:
        logger.debug("DOCX: visual hanging-indent fallback skipped: %s", e)


def _parse_span_style(style_attr: str) -> Dict:
    """Parse style='...' into color_hex, font_size_pt, font_family (for docx runs)."""
    out = {}
    if not style_attr or not isinstance(style_attr, str):
        return out
    # color: #RRGGBB or color:#RRGGBB
    m = re.search(r"color\s*:\s*#([0-9a-fA-F]{3,8})\b", style_attr)
    if m:
        hex_val = m.group(1)
        if len(hex_val) == 3:
            hex_val = "".join(c * 2 for c in hex_val)
        if len(hex_val) >= 6:
            out["color_hex"] = "#" + hex_val[:6]
    # font-size: 14pt or font-size:14pt
    m = re.search(r"font-size\s*:\s*(\d+(?:\.\d+)?)\s*pt\b", style_attr, re.IGNORECASE)
    if m:
        try:
            out["font_size_pt"] = float(m.group(1))
        except (TypeError, ValueError):
            pass
    # font-family: 'Arial' or font-family: Arial, sans-serif
    m = re.search(r"font-family\s*:\s*['\"]?([^'\";,}]+)", style_attr, re.IGNORECASE)
    if m:
        out["font_family"] = m.group(1).strip().strip("'\"").split(",")[0].strip()
    return out


class _HTMLToBlocksParser(HTMLParser):
    """Parse HTML into blocks; each run can have text, bold, italic, color, font_size_pt, font_family."""

    def __init__(self):
        super().__init__()
        self.blocks: List[Dict] = []
        self._current_block_type = "p"
        self._current_runs: List[Dict] = []
        self._bold = 0
        self._italic = 0
        self._style_stack: List[Dict] = []  # stack of {color_hex, font_size_pt, font_family} from <span>

    def _current_style(self) -> Dict:
        if not self._style_stack:
            return {}
        return dict(self._style_stack[-1])

    def _flush_run(self, text: str):
        if not text:
            return
        # Pure: no paragraph splitting from newlines; only \n → line break within current block
        parts = (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
        for i, part in enumerate(parts):
            run = {"text": part, "bold": self._bold > 0, "italic": self._italic > 0}
            run.update(self._current_style())
            self._current_runs.append(run)
            if i < len(parts) - 1:
                self._current_runs.append({"line_break": True})

    def _emit_block(self):
        if not self._current_runs:
            return
        self.blocks.append({
            "type": self._current_block_type,
            "runs": self._current_runs,
        })
        self._current_runs = []
        self._current_block_type = "p"

    def handle_starttag(self, tag: str, attrs: list):
        tag = tag.lower()
        if tag in ("p", "div", "br"):
            if tag == "br":
                # Double <br/> = paragraph break (new <w:p>); single <br/> = line break
                if self._current_runs and self._current_runs[-1].get("line_break"):
                    self._emit_block()
                else:
                    self._current_runs.append({"line_break": True})
            else:
                self._emit_block()
        elif tag in ("strong", "b"):
            self._bold += 1
        elif tag in ("em", "i"):
            self._italic += 1
        elif tag == "span":
            style_dict = {}
            for k, v in attrs:
                if k and k.lower() == "style" and v:
                    style_dict = _parse_span_style(v)
                    break
            self._style_stack.append(style_dict)
        elif tag == "li":
            self._emit_block()
            self._current_block_type = "li"
        elif tag in ("ul", "ol"):
            pass

    def handle_endtag(self, tag: str):
        tag = tag.lower()
        if tag in ("p", "div"):
            self._emit_block()
        elif tag in ("strong", "b"):
            self._bold = max(0, self._bold - 1)
        elif tag in ("em", "i"):
            self._italic = max(0, self._italic - 1)
        elif tag == "span":
            if self._style_stack:
                self._style_stack.pop()
        elif tag == "li":
            self._emit_block()
            self._current_block_type = "p"
        elif tag in ("ul", "ol"):
            pass

    def handle_data(self, data: str):
        if not data:
            return
        data = htmllib.unescape(data)
        # \n\n = new paragraph (emit block); \n = line break within paragraph
        segments = data.replace("\r\n", "\n").replace("\r", "\n").split("\n\n")
        for i, seg in enumerate(segments):
            if i > 0:
                self._emit_block()
            self._flush_run(seg)

    def get_blocks(self) -> List[Dict]:
        self._emit_block()
        return self.blocks


def _html_to_blocks(html: str) -> List[Dict]:
    """Convert HTML to a list of blocks with runs (text, bold, italic, line_break)."""
    if not html or not html.strip():
        return []
    # Keep <br/> so the parser can emit line_break runs
    parser = _HTMLToBlocksParser()
    try:
        parser.feed(html)
        return parser.get_blocks()
    except Exception as e:
        logger.warning("HTML parse failed, falling back to plain paragraphs: %s", e)
        return [{"type": "p", "runs": [{"text": re.sub(r"<[^>]+>", "", html), "bold": False, "italic": False}]}]


def _html_to_plain_paragraphs(html: str) -> list:
    """Convert HTML to a list of paragraph texts (one string per paragraph). Fallback when no formatting needed."""
    if not html or not html.strip():
        return [""]
    text = re.sub(r"</p>\s*<p(\s[^>]*)?>", "\n\n", html, flags=re.IGNORECASE)
    text = re.sub(r"<br\s*/?\s*>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</(div|h[1-6]|li|tr)>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    text = text.replace("&nbsp;", " ").replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    text = re.sub(r" +", " ", text)
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    return paragraphs if paragraphs else [""]


# def insert_line_breaks_in_long_paragraphs(text: str, max_chars: int = 80) -> str:
#     """
#     Insert single newlines in long paragraphs so the .docx gets line breaks at sentence
#     boundaries instead of one dense block. Only affects paragraphs that have no internal
#     newlines and exceed max_chars. Preserves existing \\n and \\n\\n.
#     """
#     if not text or not text.strip():
#         return text
#     normalized = (text or "").replace("\r\n", "\n\n").replace("\r", "\n\n")
#     paragraphs = re.split(r"\n\s*\n", normalized)
#     out = []
#     for para in paragraphs:
#         para = para.strip()
#         if not para:
#             out.append("")
#             continue
#         # Already has internal line breaks, or short enough
#         if "\n" in para or len(para) <= max_chars:
#             out.append(para)
#             continue
#         # Break at sentence boundaries: . " or .) " or ? " or ?) " when next char is A-Z
#         result_lines = []
#         rest = para
#         while len(rest) > max_chars:
#             chunk = rest[: max_chars + 1]
#             best_pos = -1
#             best_len = 0
#             for ending in [".) ", ". ", "?) ", "? "]:
#                 pos = chunk.rfind(ending)
#                 if pos != -1:
#                     next_idx = pos + len(ending)
#                     if next_idx < len(rest) and rest[next_idx].isalpha() and rest[next_idx].isupper():
#                         if pos > best_pos:
#                             best_pos = pos
#                             best_len = len(ending)
#             if best_pos == -1:
#                 # No sentence end; break at last space
#                 pos = chunk.rfind(" ")
#                 if pos != -1:
#                     best_pos = pos
#                     best_len = 1
#             if best_pos == -1:
#                 result_lines.append(rest)
#                 rest = ""
#                 break
#             end = best_pos + best_len
#             result_lines.append(rest[:end].rstrip())
#             rest = rest[end:].lstrip()
#         if rest:
#             result_lines.append(rest)
#         out.append("\n".join(result_lines))
#     return "\n\n".join(out)


def _strip_html_from_plain(text: str) -> str:
    """Remove HTML tags and decode entities from plain-text content so the docx never shows raw HTML."""
    if not text:
        return ""
    s = re.sub(r"<[^>]+>", "", text)
    s = htmllib.unescape(s)
    return s


def _ensure_paragraph_breaks(text: str) -> str:
    """
    Ensure we have \\n\\n at paragraph boundaries so the docx gets multiple <w:p>.
    If content already has \\n\\n, keep it. Otherwise insert \\n\\n at every sentence
    boundary: . ? ! followed by (newline or space) and then a capital letter.
    Avoids splitting after abbreviations (Mr., Dr.) by requiring 3+ chars before .?!
    for inline boundaries.
    """
    if not text:
        return ""
    normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    if "\n\n" in normalized:
        return normalized
    # 1) .?! then newline(s) then capital — likely real paragraph break
    normalized = re.sub(r"([.?!])\s*\n+\s*([A-Z])", r"\1\n\n\2", normalized)
    if "\n\n" in normalized:
        return normalized
    # 2) Inline: .?! then space(s) or no space, then capital. Require 3+ chars before .?!
    # so we don't split "Mr. Smith" or "Dr. Jones"
    normalized = re.sub(r"(?<=.{3})([.?!])\s{2,}([A-Z])", r"\1\n\n\2", normalized)
    normalized = re.sub(r"(?<=.{3})([.?!])\s+([A-Z])", r"\1\n\n\2", normalized)
    normalized = re.sub(r"(?<=.{3})([.?!])([A-Z])", r"\1\n\n\2", normalized)
    return normalized


def _strip_style_tags_from_plain(text: str) -> str:
    """
    Remove only the tag delimiters [tag:value] and [/tag], leaving all content in place.
    So tags never appear in output and we never remove or replace content.
    """
    if not text:
        return text
    # Remove opening tags: [color:#x], [ size : 14pt ], [font: Arial]
    text = re.sub(
        r"\[\s*(?:color|size|font)\s*:\s*[^\]]+\]",
        "",
        text,
        flags=re.IGNORECASE,
    )
    # Remove closing tags: [/color], [ / size ]
    text = re.sub(
        r"\[\s*/\s*(?:color|size|font)\s*\]",
        "",
        text,
        flags=re.IGNORECASE,
    )
    return text


def _strip_plain_text_formatting(line: str) -> str:
    """
    Lightweight cleanup for plain-text content:
    - remove [font:], [size:], [color:] style tags
    - strip basic markdown emphasis markers (**bold**, *italic*, __bold__, _italic_)
    without doing heavy segmentation.
    """
    if not line:
        return ""
    s = _strip_style_tags_from_plain(line)
    s = _normalize_spaced_markdown_emphasis(s)
    # Remove markdown-style emphasis while keeping inner text
    for pattern in (
        r"\*\*(.+?)\*\*",  # **bold**
        r"__(.+?)__",      # __bold__
        r"\*(.+?)\*",      # *italic*
        r"_(.+?)_",        # _italic_
    ):
        s = re.sub(pattern, r"\1", s)
    return s.strip()


def _normalize_spaced_markdown_emphasis(text: str) -> str:
    """
    Normalize loose LLM emphasis markers like "* * text * *" into "**text**".
    Supports both asterisk and underscore variants.
    """
    if not text:
        return ""
    s = text
    # Convert "* * text * *" -> "**text**"
    s = re.sub(r"(?<!\*)\*\s+\*(.+?)\*\s+\*(?!\*)", r"**\1**", s)
    # Convert "_ _ text _ _" -> "__text__"
    s = re.sub(r"(?<!_)_\s+_(.+?)_\s+_(?!_)", r"__\1__", s)
    return s


# Max length of a single tag [...] we will parse; beyond this we treat as literal (no hang).
_MAX_TAG_LEN = 120


def _parse_style_segments(line: str) -> List[tuple]:
    """
    Split line by [color:#xxx], [size:Npt], [font:Name] and [/color], [/size], [/font].
    Uses only str.find() and short slices — no regex on unbounded input. Safe and linear.
    """
    segments = []
    style_stack: List[Dict] = [{}]
    pos = 0
    n = len(line)

    def current_style() -> Dict:
        merged = {}
        for d in style_stack:
            merged.update(d)
        return merged

    while pos < n:
        lb = line.find("[", pos)
        if lb == -1:
            if pos < n:
                segments.append((current_style(), line[pos:]))
            break
        if lb > pos:
            segments.append((current_style(), line[pos:lb]))
        rb = line.find("]", lb + 1)
        if rb == -1 or (rb - lb) > _MAX_TAG_LEN:
            segments.append((current_style(), line[lb : lb + 1]))
            pos = lb + 1
            continue
        tag_inner = line[lb + 1 : rb].strip().lower()
        pos = rb + 1
        if tag_inner.startswith("/"):
            closer = tag_inner[1:].strip()
            if len(style_stack) > 1 and closer in ("color", "size", "font"):
                style_stack.pop()
            continue
        if ":" in tag_inner:
            kind, _, value = tag_inner.partition(":")
            value = value.strip()
            kind = kind.strip()
            new_style = {}
            if kind == "color" and value:
                hex_val = value.lstrip("#")[:8]
                if len(hex_val) == 3:
                    hex_val = "".join(c * 2 for c in hex_val)
                if len(hex_val) >= 6:
                    new_style["color_hex"] = "#" + hex_val[:6]
            elif kind == "size" and value:
                num_str = ""
                for c in value:
                    if c.isdigit() or c == ".":
                        num_str += c
                    else:
                        break
                if num_str:
                    try:
                        new_style["font_size_pt"] = float(num_str)
                    except ValueError:
                        pass
            elif kind == "font" and value:
                new_style["font_family"] = value[:80]
            if new_style:
                style_stack.append(dict(current_style(), **new_style))
    return segments


def _find_next_bold_italic_open(text: str, start: int):
    """
    Find the earliest opening delimiter **, __, *, or _ after start.
    For * and _ we require they are not the start of ** or __.
    Returns (position, pattern_len, is_bold) or None if none found.
    """
    n = len(text)
    best = None  # (pos, pattern_len, is_bold)
    i = text.find("**", start)
    if i != -1 and (best is None or i < best[0]):
        best = (i, 2, True)
    i = text.find("__", start)
    if i != -1 and (best is None or i < best[0]):
        best = (i, 2, True)
    i = text.find("*", start)
    if i != -1 and (i + 1 >= n or text[i + 1] != "*") and (best is None or i < best[0]):
        best = (i, 1, False)  # italic
    i = text.find("_", start)
    if i != -1 and (i + 1 >= n or text[i + 1] != "_") and (best is None or i < best[0]):
        best = (i, 1, False)  # italic
    return best


def _find_bold_italic_close(text: str, open_pos: int, pattern_len: int, is_double: bool) -> int:
    """Find the closing delimiter. For double (** or __) next occurrence. For single (* or _) next occurrence not part of double."""
    n = len(text)
    search_start = open_pos + pattern_len
    if is_double:
        i = text.find(text[open_pos : open_pos + pattern_len], search_start)
        return i
    needle = text[open_pos]
    j = search_start
    while j < n:
        j = text.find(needle, j)
        if j == -1:
            return -1
        if j + 1 >= n or text[j + 1] != needle:
            return j
        j += 2
    return -1


def _bold_italic_to_runs(text: str) -> List[Dict]:
    """
    Parse **bold**, *italic*, __bold__, _italic_ with linear scan (str.find only).
    No backtracking regex, so safe from catastrophic hang on bad input.
    """
    if not text:
        return [{"text": "", "bold": False, "italic": False}]
    runs = []
    pos = 0
    n = len(text)
    while pos < n:
        best = _find_next_bold_italic_open(text, pos)
        if best is None:
            runs.append({"text": text[pos:], "bold": False, "italic": False})
            break
        open_pos, pattern_len, is_bold = best
        is_double = pattern_len == 2
        if open_pos > pos:
            runs.append({"text": text[pos:open_pos], "bold": False, "italic": False})
        close_pos = _find_bold_italic_close(text, open_pos, pattern_len, is_double)
        if close_pos == -1:
            runs.append({"text": text[open_pos : open_pos + 1], "bold": False, "italic": False})
            pos = open_pos + 1
            continue
        inner = text[open_pos + pattern_len : close_pos]
        runs.append({"text": inner, "bold": is_bold, "italic": not is_bold})
        pos = close_pos + pattern_len
    return runs if runs else [{"text": text, "bold": False, "italic": False}]


_MAX_RUNS_PER_LINE = 300  # Cap so we never build huge run lists (safety and perf)


def _plain_line_to_runs(line: str) -> List[Dict]:
    """Parse one line: [color:#x][/color], [size:Npt][/size], [font:Name][/font] plus **bold** and *italic*."""
    runs = []
    normalized_line = _normalize_spaced_markdown_emphasis(line)
    segments = _parse_style_segments(normalized_line)
    if not segments:
        return [{"text": _strip_style_tags_from_plain(line), "bold": False, "italic": False}]
    for style_dict, segment_text in segments:
        if len(runs) >= _MAX_RUNS_PER_LINE:
            break
        for run in _bold_italic_to_runs(segment_text):
            if len(runs) >= _MAX_RUNS_PER_LINE:
                break
            run = dict(run)
            if style_dict:
                if "color_hex" in style_dict:
                    run["color_hex"] = style_dict["color_hex"]
                if "font_size_pt" in style_dict:
                    run["font_size_pt"] = style_dict["font_size_pt"]
                if "font_family" in style_dict:
                    run["font_family"] = style_dict["font_family"]
            runs.append(run)
    if not runs:
        return [{"text": _strip_style_tags_from_plain(line), "bold": False, "italic": False}]
    return runs


# Max line length for full style/markdown parsing. Longer lines use stripped plain text only.
_SAFE_LINE_LENGTH = 12000


def _plain_text_to_blocks(text: str) -> List[Dict]:
    """
    \\n\\n = new paragraph (<w:p>); \\n = line break within paragraph.
    Full parsing: [font:...], [size:...], [color:...] and **bold** / *italic* (linear-time, no hang).
    Lines longer than _SAFE_LINE_LENGTH are stripped only to avoid edge cases.
    """
    if not text:
        return []
    normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    paragraphs = re.split(r"\n\s*\n", normalized.strip() if normalized else "")
    blocks = []

    # Accept common ordered-list prefixes from LLMs.
    # Examples: "1. item", "1) item", "(1) item"
    # Intentionally limited to 1-2 digits to avoid false positives on phone
    # patterns like "(818) 419-5986".
    list_number_re = re.compile(r"^\s*(?:\(?[1-9]\d?\)?[.)])\s+")
    # Accept common bullet prefixes.
    # Examples: "- item", "* item", "+ item", "• item", "◦ item", "▪ item", "▸ item"
    list_bullet_re = re.compile(r"^\s*(?:[•◦▪▸\-\*\+])\s+")

    for para in paragraphs:
        lines = para.split("\n")
        # If every non-empty line is a list item, force one DOCX paragraph per line.
        non_empty = [ln for ln in lines if ln.strip()]
        is_all_list_lines = bool(non_empty) and all(
            list_number_re.match(ln) or list_bullet_re.match(ln) for ln in non_empty
        )
        # If a paragraph mixes normal text and list lines (e.g. heading + bullets),
        # split per line so list lines can be promoted to Word-native lists later.
        has_any_list_lines = any(
            list_number_re.match(ln) or list_bullet_re.match(ln) for ln in non_empty
        )

        if is_all_list_lines or has_any_list_lines:
            for line in lines:
                if not line.strip():
                    continue
                if len(line) <= _SAFE_LINE_LENGTH:
                    runs = _plain_line_to_runs(line)
                else:
                    clean = _strip_plain_text_formatting(line)
                    runs = [{"text": clean, "bold": False, "italic": False}] if clean else []
                if runs:
                    blocks.append({"type": "p", "runs": runs})
            continue

        runs = []
        for i, line in enumerate(lines):
            if len(line) <= _SAFE_LINE_LENGTH:
                runs.extend(_plain_line_to_runs(line))
            else:
                clean = _strip_plain_text_formatting(line)
                if clean:
                    runs.append({"text": clean, "bold": False, "italic": False})
            if i < len(lines) - 1:
                runs.append({"line_break": True})
        blocks.append({"type": "p", "runs": runs})
    return blocks


def _markdown_to_blocks(markdown: str) -> List[Dict]:
    """Convert markdown to blocks with runs (preserve ** and * as bold/italic)."""
    if not markdown or not markdown.strip():
        return []
    blocks = []
    # Split into paragraphs (double newline)
    raw_paras = re.split(r"\n\s*\n", markdown.strip())
    for raw in raw_paras:
        line = " ".join(raw.splitlines()).strip()
        line = _normalize_spaced_markdown_emphasis(line)
        if not line:
            continue
        # Detect list item
        is_li = bool(re.match(r"^[\-\*]\s+", line) or re.match(r"^\d+\.\s+", line))
        line = re.sub(r"^[\-\*]\s+", "", line)
        line = re.sub(r"^\d+\.\s+", "", line)
        # Split into runs by ** and *
        runs = []
        rest = line
        while rest:
            # Bold: **text**
            m = re.search(r"\*\*(.+?)\*\*", rest)
            if m:
                before = rest[: m.start()]
                if before:
                    runs.append({"text": before, "bold": False, "italic": False})
                runs.append({"text": m.group(1), "bold": True, "italic": False})
                rest = rest[m.end() :]
                continue
            # Italic: *text*
            m = re.search(r"\*(.+?)\*", rest)
            if m:
                before = rest[: m.start()]
                if before:
                    runs.append({"text": before, "bold": False, "italic": False})
                runs.append({"text": m.group(1), "bold": False, "italic": True})
                rest = rest[m.end() :]
                continue
            # __bold__
            m = re.search(r"__(.+?)__", rest)
            if m:
                before = rest[: m.start()]
                if before:
                    runs.append({"text": before, "bold": False, "italic": False})
                runs.append({"text": m.group(1), "bold": True, "italic": False})
                rest = rest[m.end() :]
                continue
            # _italic_
            m = re.search(r"_(.+?)_", rest)
            if m:
                before = rest[: m.start()]
                if before:
                    runs.append({"text": before, "bold": False, "italic": False})
                runs.append({"text": m.group(1), "bold": False, "italic": True})
                rest = rest[m.end() :]
                continue
            runs.append({"text": rest, "bold": False, "italic": False})
            break
        if not runs:
            runs = [{"text": line, "bold": False, "italic": False}]
        blocks.append({"type": "li" if is_li else "p", "runs": runs})
    return blocks


def _set_paragraph_num_pr(paragraph, ilvl: int, num_id: int) -> None:
    """Set w:numPr (ilvl, numId) on a paragraph so Word renders it as a list item."""
    if not DOCX_AVAILABLE or OxmlElement is None or qn is None:
        return
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


def _ensure_docx_list_numbering(doc) -> tuple:
    """
    Ensure the document has bullet (numId 100) and decimal (numId 101) list definitions
    in word/numbering.xml per OOXML. Returns (bullet_num_id, number_num_id) or (None, None).
    Build elements with OxmlElement to avoid parse_xml attribute escaping issues.
    """
    if not DOCX_AVAILABLE or OxmlElement is None or qn is None:
        return (None, None)
    try:
        numbering_part = doc.part.numbering_part
    except AttributeError:
        logger.debug("DOCX: no numbering part on document")
        return (None, None)
    try:
        root = numbering_part.element
    except AttributeError:
        root = getattr(
            getattr(numbering_part, "numbering_definitions", None), "_numbering", None
        )
    if root is None:
        return (None, None)
    # Use high IDs to avoid clashing with template's built-in numbering
    BULLET_ABSTRACT_ID, BULLET_NUM_ID = 100, 100
    NUMBER_ABSTRACT_ID, NUMBER_NUM_ID = 101, 101

    def _el(tag: str, attrs: Optional[Dict[str, str]] = None):
        e = OxmlElement(tag)
        if attrs:
            for k, v in attrs.items():
                e.set(qn(k), str(v))
        return e

    def _abstract_num(abstract_num_id: int, num_fmt: str, lvl_text: str):
        abstract = _el("w:abstractNum", {"w:abstractNumId": str(abstract_num_id)})
        abstract.append(_el("w:multiLevelType", {"w:val": "singleLevel"}))
        lvl = _el("w:lvl", {"w:ilvl": "0"})
        lvl.append(_el("w:numFmt", {"w:val": num_fmt}))
        lvl.append(_el("w:lvlText", {"w:val": lvl_text}))
        lvl.append(_el("w:lvlJc", {"w:val": "left"}))
        pPr = OxmlElement("w:pPr")
        pPr.append(_el("w:ind", {"w:left": "720", "w:hanging": "360"}))
        lvl.append(pPr)
        rPr = OxmlElement("w:rPr")
        rPr.append(_el("w:rFonts", {"w:ascii": "Symbol", "w:hAnsi": "Symbol"}))
        lvl.append(rPr)
        abstract.append(lvl)
        return abstract

    def _num(num_id: int, abstract_num_id: int):
        num = _el("w:num", {"w:numId": str(num_id)})
        num.append(_el("w:abstractNumId", {"w:val": str(abstract_num_id)}))
        return num

    bullet_abstract = _abstract_num(BULLET_ABSTRACT_ID, "bullet", "\u2022")
    bullet_num = _num(BULLET_NUM_ID, BULLET_ABSTRACT_ID)
    number_abstract = _abstract_num(NUMBER_ABSTRACT_ID, "decimal", "%1.")
    number_num = _num(NUMBER_NUM_ID, NUMBER_ABSTRACT_ID)
    root.append(bullet_abstract)
    root.append(bullet_num)
    root.append(number_abstract)
    root.append(number_num)
    return (BULLET_NUM_ID, NUMBER_NUM_ID)


def build_docx_from_content(
    content: str,
    *,
    from_html: bool = False,
    from_plain_text: bool = False,
    print_properties: Optional[Dict] = None,
) -> bytes:
    """
    Build a Word .docx from cover letter content.

    Args:
        content: The cover letter body (plain text, markdown, or HTML).
        from_html: If True, treat content as HTML.
        from_plain_text: If True, treat content as plain text (\\n = line break, \\n\\n = paragraph). No markup.
        print_properties: Optional dict with fontFamily, fontSize, lineHeight.

    Returns:
        .docx file as bytes.
    """
    logger.info(
        "DOCX: build_docx_from_content start (from_html=%s, from_plain_text=%s)",
        from_html,
        from_plain_text,
    )
    if not DOCX_AVAILABLE or Document is None:
        logger.error("DOCX: python-docx not available; cannot build .docx")
        raise ImportError("python-docx is not installed. Install with: pip install python-docx")

    props = print_properties or {}
    use_default_fonts = props.get("useDefaultFonts", False)
    font_family = props.get("fontFamily", "Times New Roman")
    font_size_pt = props.get("fontSize", 12)
    try:
        font_size_pt = float(font_size_pt)
    except (TypeError, ValueError):
        font_size_pt = 12
    line_height = props.get("lineHeight", 1.6)
    try:
        line_height = float(line_height)
    except (TypeError, ValueError):
        line_height = 1.6

    if from_plain_text:
        content = _strip_html_from_plain(content or "")
        blocks = _plain_text_to_blocks(content)
        logger.info("DOCX: using plain-text path; content length=%s", len(content or ""))
    elif from_html:
        logger.info("DOCX: using HTML path; content length=%s", len(content or ""))
        blocks = _html_to_blocks(content)
    else:
        logger.info("DOCX: using markdown path; content length=%s", len(content or ""))
        blocks = _markdown_to_blocks(content or "")

    logger.info("DOCX: blocks parsed (count=%s)", len(blocks))

    # Strip any [font:...], [size:...], [color:...] and [/font], [/size], [/color] that ended up as literal run text
    # (e.g. parser edge cases, or when content came from a path that doesn't parse these tags). Keeps docx clean.
    # Also strip any raw ** or __ that slipped through (e.g. "**             **" when inner is whitespace-only).
    for block in blocks:
        for run in block.get("runs", []):
            if "text" in run and run["text"]:
                t = run["text"]
                while True:
                    t2 = _strip_style_tags_from_plain(t)
                    if t2 == t:
                        break
                    t = t2
                t = t.replace("**", "").replace("__", "")
                run["text"] = t

    # Convert bullet/number prefixes to Word native list styles (List Bullet / List Number)
    # Markdown bullets: - * + (with space after). Unicode: • ◦ ▪ ▸ U+2022
    _BULLET_CHARS = ("\u2022", "•", "-", "*", "+", "\u25E6", "\u25AA", "\u25B8")
    _BULLET_REQUIRE_SPACE = ("-", "*", "+")  # require space/tab after so "non-profit", "5*4", "C++" aren't bullets
    for block in blocks:
        if block.get("type") != "p" or not block.get("runs"):
            continue
        # Use first run that has non-empty text (bullet may not be in runs[0] if paragraph starts with blank line)
        first_run = None
        for r in block["runs"]:
            if r.get("line_break"):
                continue
            t = (r.get("text") or "").strip()
            if t:
                first_run = r
                break
        if first_run is None:
            continue
        first_text = first_run.get("text", "") or ""
        stripped = first_text.lstrip()
        if not stripped:
            continue
        # Numbered list: "1. ", "2. ", "10. " etc. -> Word List Number
        # Limit to 1-2 digits so contact lines like "(818) ..." are not treated as lists.
        num_match = re.match(r"^(?:\(?([1-9]\d?)\)?[.)])\s+", stripped)
        if num_match:
            block["type"] = "li_number"
            first_run["text"] = stripped[num_match.end() :]
            # Remove leading line-break/empty runs so list item doesn't start with a blank line
            while block["runs"] and (block["runs"][0].get("line_break") or not (block["runs"][0].get("text") or "").strip()):
                block["runs"].pop(0)
            continue
        # Bullet list: bullet char + optional whitespace -> Word List Bullet
        for bullet in _BULLET_CHARS:
            if not stripped.startswith(bullet):
                continue
            if bullet in _BULLET_REQUIRE_SPACE and len(stripped) > len(bullet) and stripped[len(bullet)] not in " \t":
                continue
            rest = stripped[len(bullet) :].lstrip()
            block["type"] = "li"
            first_run["text"] = rest
            # Remove leading line-break/empty runs so list item doesn't start with a blank line
            while block["runs"] and (block["runs"][0].get("line_break") or not (block["runs"][0].get("text") or "").strip()):
                block["runs"].pop(0)
            break

    logger.info("DOCX: creating Document and applying styles")
    doc = Document()
    bullet_num_id, number_num_id = _ensure_docx_list_numbering(doc)
    use_numbering_part = bullet_num_id is not None and number_num_id is not None
    # Apply section margins and page size from print_properties (client sends margins in inches)
    if DOCX_AVAILABLE and Inches is not None:
        section = doc.sections[0]
        margins = props.get("margins")
        if isinstance(margins, dict):
            for key, attr in (("top", "top_margin"), ("right", "right_margin"), ("bottom", "bottom_margin"), ("left", "left_margin")):
                val = margins.get(key)
                if val is not None:
                    try:
                        setattr(section, attr, Inches(float(val)))
                    except (TypeError, ValueError):
                        pass
        page_size = props.get("pageSize")
        if isinstance(page_size, dict):
            w = page_size.get("width")
            h = page_size.get("height")
            if w is not None:
                try:
                    section.page_width = Inches(float(w))
                except (TypeError, ValueError):
                    pass
            if h is not None:
                try:
                    section.page_height = Inches(float(h))
                except (TypeError, ValueError):
                    pass
    style = doc.styles["Normal"]
    # Always apply user's line height (line spacing)
    style.paragraph_format.line_spacing = line_height
    # Apply font/size only when user has not chosen default fonts
    if not use_default_fonts and font_family and str(font_family).strip().lower() != "default":
        style.font.name = font_family
        style.font.size = Pt(font_size_pt)
    else:
        font_family = font_family or "Times New Roman"
        font_size_pt = font_size_pt if font_size_pt else 12

    space_after = Pt(round(font_size_pt * line_height * 0.4))

    # List formatting: either OOXML numbering (numPr + numbering.xml) or fallback (indent + text bullet/number).
    _LIST_LEFT = Inches(0.25)
    _LIST_HANG = Inches(-0.25)

    # One <w:p> per block (see BACKEND_DOCX_PARAGRAPH_AND_LINE_BREAKS.md)
    logger.info("DOCX: starting paragraph/run construction")
    list_number = 0
    for block in blocks:
        p = doc.add_paragraph()
        if block["type"] == "li":
            if use_numbering_part:
                _set_paragraph_num_pr(p, 0, bullet_num_id)
            else:
                p.paragraph_format.left_indent = _LIST_LEFT
                p.paragraph_format.first_line_indent = _LIST_HANG
        elif block["type"] == "li_number":
            list_number += 1
            if use_numbering_part:
                _set_paragraph_num_pr(p, 0, number_num_id)
            else:
                p.paragraph_format.left_indent = _LIST_LEFT
                p.paragraph_format.first_line_indent = _LIST_HANG
        else:
            list_number = 0
            p.paragraph_format.space_after = space_after
        last_ended_with_space = True  # avoid leading space before first run
        first_text_run = True  # prepend bullet/number only when not using numbering part
        for run_spec in block["runs"]:
            if run_spec.get("line_break"):
                if WD_BREAK is not None:
                    p.add_run().add_break(WD_BREAK.LINE)
                last_ended_with_space = True
                continue
            text = run_spec.get("text", "")
            if not text:
                continue
            # Prepend visible bullet/number only when not using OOXML numbering
            if first_text_run and block["type"] in ("li", "li_number") and not use_numbering_part:
                first_text_run = False
                if block["type"] == "li":
                    text = "\u2022 " + text
                else:
                    text = f"{list_number}. " + text
            # Preserve space around inline/spans: add space if previous run didn't end with space and this doesn't start with one
            if not last_ended_with_space and not (text.startswith(" ") or text.startswith("\t")):
                text = " " + text
            last_ended_with_space = text.endswith(" ") or text.endswith("\t")
            run = p.add_run(text)
            run_font = run.font
            inline_font_family = (run_spec.get("font_family") or "").strip()
            inline_size_pt = run_spec.get("font_size_pt")
            has_inline_size = inline_size_pt is not None and inline_size_pt > 0

            # Always honor explicit inline style tags from content (e.g. [size:28pt], [font:Georgia]),
            # even when useDefaultFonts is enabled.
            if inline_font_family:
                run_font.name = inline_font_family
            elif not use_default_fonts:
                run_font.name = font_family

            if has_inline_size:
                run_font.size = Pt(inline_size_pt)
            elif not use_default_fonts:
                run_font.size = Pt(font_size_pt)
            run.bold = run_spec.get("bold", False)
            run.italic = run_spec.get("italic", False)
            color_hex = run_spec.get("color_hex")
            if color_hex and RGBColor is not None:
                hex_val = color_hex.lstrip("#")
                if len(hex_val) >= 6:
                    try:
                        r = int(hex_val[0:2], 16)
                        g = int(hex_val[2:4], 16)
                        b = int(hex_val[4:6], 16)
                        run_font.color.rgb = RGBColor(r, g, b)
                    except (ValueError, TypeError):
                        pass

    logger.info("DOCX: writing document to bytes buffer")
    _apply_visual_hanging_indent_fallback(doc)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.read()
    logger.info("DOCX: document built (size=%s bytes)", len(docx_bytes))

    # Debug: always write docx + info to tmp/ after every build (so we see which path ran: plain_text vs html vs markdown).
    _source = "plain_text" if from_plain_text else ("html" if from_html else "markdown")
    _written = []
    _bases = [
        os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..")),
        os.getcwd(),
    ]
    for _base in _bases:
        _tmp = os.path.join(_base, "tmp")
        _docx_path = os.path.join(_tmp, "raw-debug.docx")
        _info_path = os.path.join(_tmp, "raw-debug-info.txt")
        try:
            os.makedirs(_tmp, exist_ok=True)
            with open(_docx_path, "wb") as f:
                f.write(docx_bytes)
            with open(_info_path, "w", encoding="utf-8") as f:
                f.write(f"source={_source}\n")
                f.write(f"block_count={len(blocks)}\n")
                raw = (content or "")[:1200]
                has_double = "\n\n" in raw
                f.write(f"content_has_double_newline={has_double}\n")
                f.write(f"content_preview (repr)=\n{repr(raw)}\n")
            _written.append(os.path.abspath(_tmp))
        except Exception as e:
            logger.warning("Docx debug: could not write to %s: %s", os.path.abspath(_tmp), e)
    if _written:
        logger.info(
            "Docx debug: wrote raw-debug.docx, raw-debug-info.txt (source=%s, blocks=%s) to %s",
            _source,
            len(blocks),
            _written,
        )
    else:
        logger.warning("Docx debug: no tmp dir was writable")

    logger.info("DOCX: build_docx_from_content finished successfully")
    return docx_bytes


# Minimal OOXML skeletons for build_docx_from_components when LLM omits numbering or styles
_MINIMAL_CONTENT_TYPES = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/numbering.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>"""

_MINIMAL_RELS = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""

_MINIMAL_DOCUMENT_RELS = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering" Target="numbering.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>
</Relationships>"""

_MINIMAL_NUMBERING = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:abstractNum w:abstractNumId="0">
    <w:multiLevelType w:val="singleLevel"/>
    <w:lvl w:ilvl="0">
      <w:numFmt w:val="bullet"/>
      <w:lvlText w:val="&#x2022;"/>
      <w:lvlJc w:val="left"/>
      <w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>
      <w:rPr><w:rFonts w:ascii="Symbol" w:hAnsi="Symbol"/></w:rPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>
</w:numbering>"""

_MINIMAL_STYLES = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:docDefaults><w:rPrDefault/><w:pPrDefault/></w:docDefaults>
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
    <w:name w:val="Normal"/>
    <w:qFormat/>
  </w:style>
</w:styles>"""


def build_docx_from_components(
    document_xml: str,
    numbering_xml: Optional[str] = None,
    styles_xml: Optional[str] = None,
) -> bytes:
    """
    Assemble a .docx (ZIP) from the three Word OOXML components returned by the LLM.

    Args:
        document_xml: Full content of word/document.xml (root <w:document> with <w:body> and <w:sectPr/>).
        numbering_xml: Optional content of word/numbering.xml; if missing, minimal bullet numbering is used.
        styles_xml: Optional content of word/styles.xml; if missing, minimal styles are used.

    Returns:
        .docx file as bytes.
    """
    doc_xml = (document_xml or "").strip()
    if not doc_xml:
        raise ValueError("document_xml is required and must be non-empty")
    num_xml = (numbering_xml or "").strip() or _MINIMAL_NUMBERING
    sty_xml = (styles_xml or "").strip() or _MINIMAL_STYLES

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", _MINIMAL_CONTENT_TYPES)
        zf.writestr("_rels/.rels", _MINIMAL_RELS)
        zf.writestr("word/_rels/document.xml.rels", _MINIMAL_DOCUMENT_RELS)
        zf.writestr("word/document.xml", doc_xml)
        zf.writestr("word/numbering.xml", num_xml)
        zf.writestr("word/styles.xml", sty_xml)
    buf.seek(0)
    out = buf.read()
    logger.info("DOCX: assembled from components (document=%s bytes, numbering=%s, styles=%s)", len(doc_xml), bool(numbering_xml), bool(styles_xml))
    return out


def apply_print_properties_to_docx(
    docx_bytes: bytes,
    print_properties: Optional[Dict] = None,
) -> bytes:
    """
    Apply user print settings to an existing .docx (e.g. one built from components).
    - Line height: always applied to Normal style.
    - Font family/size: applied only when useDefaultFonts is False (same rule as elsewhere).
    """
    if not DOCX_AVAILABLE or Document is None or Pt is None:
        return docx_bytes
    props = print_properties or {}
    use_default_fonts = props.get("useDefaultFonts", False)
    line_height = props.get("lineHeight", 1.6)
    try:
        line_height = float(line_height)
    except (TypeError, ValueError):
        line_height = 1.6
    font_family = props.get("fontFamily", "Times New Roman")
    font_size_pt = props.get("fontSize", 12)
    try:
        font_size_pt = float(font_size_pt)
    except (TypeError, ValueError):
        font_size_pt = 12

    try:
        doc = Document(io.BytesIO(docx_bytes))
        # Apply section margins and page size so component-built docs honor user settings on first open.
        if Inches is not None and doc.sections:
            section = doc.sections[0]
            margins = props.get("margins")
            if isinstance(margins, dict):
                for key, attr in (
                    ("top", "top_margin"),
                    ("right", "right_margin"),
                    ("bottom", "bottom_margin"),
                    ("left", "left_margin"),
                ):
                    val = margins.get(key)
                    if val is not None:
                        try:
                            setattr(section, attr, Inches(float(val)))
                        except (TypeError, ValueError):
                            pass
            page_size = props.get("pageSize")
            if isinstance(page_size, dict):
                w = page_size.get("width")
                h = page_size.get("height")
                if w is not None:
                    try:
                        section.page_width = Inches(float(w))
                    except (TypeError, ValueError):
                        pass
                if h is not None:
                    try:
                        section.page_height = Inches(float(h))
                    except (TypeError, ValueError):
                        pass
        style = doc.styles["Normal"]
        # Always apply line spacing (user's line height setting)
        style.paragraph_format.line_spacing = line_height
        # Apply font only when user has not chosen "default fonts"
        if not use_default_fonts and font_family and str(font_family).strip().lower() != "default":
            style.font.name = (font_family or "Times New Roman").strip()
            style.font.size = Pt(font_size_pt)
        _apply_visual_hanging_indent_fallback(doc)
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        logger.info("DOCX: applied print_properties (line_height=%.2f, font=%s)", line_height, "default" if use_default_fonts else font_family)
        return out.read()
    except Exception as e:
        logger.warning("DOCX: could not apply print_properties to docx: %s", e)
        return docx_bytes


def build_docx_from_generation_result(
    content: Optional[str] = None,
    markdown: Optional[str] = None,
    html: Optional[str] = None,
    print_properties: Optional[Dict] = None,
    *,
    use_plain_text: bool = False,
) -> bytes:
    """
    Build .docx from cover letter generation result.
    When use_plain_text is True (docx-only flow), content is plain text only.
    Otherwise: prefers content as plain text if use_plain_text; else prefers html, then markdown.
    """
    logger.info(
        "DOCX: build_docx_from_generation_result start (use_plain_text=%s, has_content=%s, has_markdown=%s, has_html=%s)",
        use_plain_text,
        bool(content),
        bool(markdown),
        bool(html and html.strip()),
    )
    if use_plain_text and content:
        logger.info("DOCX: generation_result using plain-text content path")
        return build_docx_from_content(
            content, from_plain_text=True, print_properties=print_properties
        )
    if html and html.strip():
        logger.info("DOCX: generation_result using HTML path")
        return build_docx_from_content(html, from_html=True, print_properties=print_properties)
    raw = content or markdown or ""
    if use_plain_text:
        logger.info("DOCX: generation_result using raw plain-text path (no HTML/markdown)")
        return build_docx_from_content(raw, from_plain_text=True, print_properties=print_properties)
    logger.info("DOCX: generation_result using markdown/plain fallback path")
    return build_docx_from_content(raw, from_html=False, print_properties=print_properties)
